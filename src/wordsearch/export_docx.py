from docx import Document


def save_wordsearch_to_docx(grid, words, filename):
    """
    Saves the wordsearch result to a DOCX file.

    Args:
        wordsearch_result (dict): The result from wordsearch.generate,
                                  expected to have 'grid' and 'words' keys.
        filename (str): The path to the output DOCX file.
    """
    doc = Document()
    doc.add_heading("Word Search Puzzle", 0)

    # Add the grid
    doc.add_heading("Puzzle Grid", level=1)
    table = doc.add_table(rows=len(grid), cols=len(grid[0]))
    for i, row in enumerate(grid):
        for j, letter in enumerate(row):
            table.cell(i, j).text = letter

    # Add the word list
    doc.add_heading("Word List", level=1)
    for word in words:
        doc.add_paragraph(word, style="List Bullet")

    doc.save(filename)


if __name__ == "__main__":
    # Example usage
    wordsearch_result = {
        "grid": [
            ["A", "B", "C", "D"],
            ["E", "F", "G", "H"],
            ["I", "J", "K", "L"],
            ["M", "N", "O", "P"],
        ],
        "words": ["ABCD", "EFGH", "IJKL", "MNOP"],
    }
    save_wordsearch_to_docx(
        wordsearch_result["grid"], wordsearch_result["words"], "wordsearch.docx"
    )
