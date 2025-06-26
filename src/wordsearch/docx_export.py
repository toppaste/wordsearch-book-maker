from docx import Document
import os

def save_wordsearch_to_docx(output_path, title, grid, words, solution):
    """
    Saves the wordsearch result to a DOCX file.

    Args:
        output_path (str): The path to the output DOCX file.
        title (str): The title of the word search puzzle.
        grid (list): The grid of letters for the puzzle.
        words (list): The list of words to find.
        solution (list): The solution for the puzzle.
    """
    doc = Document()
    doc.add_heading(title, 0)

    # Add the grid
    doc.add_heading("Puzzle Grid", level=1)
    table = doc.add_table(rows=len(grid), cols=len(grid[0]))
    for i, row in enumerate(grid):
        for j, letter in enumerate(row):
            table.cell(i, j).text = letter

    # Add the word list
    doc.add_heading("Word List", level=1)
    for word in words:
        doc.add_paragraph(word, style="List Bullet")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    # Example usage
    wordsearch_result = {
        "grid": [
            ["A", "B", "C", "D"],
            ["E", "F", "G", "H"],
            ["I", "J", "K", "L"],
            ["M", "N", "O", "P"],
        ],
        "words": ["ABCD", "EFGH", "IJKL", "MNOP"],
    }
    output_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "out", "wordsearch.docx"))
    save_wordsearch_to_docx(
        output_path,
        "title",
        wordsearch_result["grid"],
        wordsearch_result["words"],
        None
    )
